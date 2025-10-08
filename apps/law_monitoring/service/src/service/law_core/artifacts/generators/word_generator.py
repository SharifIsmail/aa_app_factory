"""
Word document generator for law summary reports matching HTML and PDF structure.
"""

import re
from datetime import datetime
from io import BytesIO

from docx import Document as create_document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from service.law_core.models import LawSummaryData
from service.models import OfficialJournalSeries

from .base_generator import BaseReportGenerator


class WordReportGenerator(BaseReportGenerator):
    """
    Word document report generator for law summary reports.

    Creates Word documents that match the structure and layout of HTML and PDF versions,
    with identical data presentation and professional formatting.
    """

    def __init__(self) -> None:
        """Initialize the Word generator."""
        super().__init__()

    def render(self, law_summary_data: LawSummaryData) -> bytes:
        """
        Render law summary data as Word document matching HTML/PDF structure exactly.

        Args:
            law_summary_data: Structured law summary analysis data

        Returns:
            bytes: Word document content
        """
        doc = create_document()

        # Set up page margins
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Setup professional styles matching PDF
        self._setup_styles(doc)

        # REPORT HEADER SECTION
        title_text = law_summary_data.title
        title = doc.add_paragraph(title_text, style="MainTitle")

        # Published date
        if law_summary_data.publication_date:
            doc.add_paragraph(
                f"Published: {law_summary_data.publication_date}", style="Metadata"
            )

        # Generated at metadata
        generated_at = law_summary_data.metadata.get(
            "generated_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")
        )
        doc.add_paragraph(f"Generated at: {generated_at}", style="Metadata")

        # Law ID metadata
        if law_summary_data.law_id:
            doc.add_paragraph(f"Law ID: {law_summary_data.law_id}", style="Metadata")

        doc.add_paragraph()  # Spacing

        # BASIC INFORMATION SECTION - Now includes summary content (matching HTML template structure)
        doc.add_paragraph("Basic Information", style="SectionHeader")

        if law_summary_data.document_date:
            self._add_info_item(
                doc, "DOCUMENT DATE", str(law_summary_data.document_date)
            )
        else:
            self._add_info_item(doc, "DOCUMENT DATE", "")

        if law_summary_data.publication_date:
            self._add_info_item(
                doc, "PUBLICATION DATE", str(law_summary_data.publication_date)
            )
        else:
            self._add_info_item(doc, "PUBLICATION DATE", "Not explicitly stated")

        if law_summary_data.date_of_effect:
            self._add_info_item(
                doc, "DATE OF EFFECT", str(law_summary_data.date_of_effect)
            )
        else:
            self._add_info_item(doc, "DATE OF EFFECT", "Not explicitly stated")

        if law_summary_data.end_validity_date:
            self._add_info_item(
                doc, "END OF VALIDITY", str(law_summary_data.end_validity_date)
            )
        else:
            self._add_info_item(doc, "END OF VALIDITY", "Not explicitly stated")

        if law_summary_data.document_type_label:
            self._add_info_item(
                doc, "DOCUMENT TYPE", str(law_summary_data.document_type_label)
            )

        oj_series_label = law_summary_data.oj_series_label
        if oj_series_label and str(oj_series_label) != OfficialJournalSeries.UNKNOWN:
            self._add_info_item(doc, "JOURNAL SERIES", str(oj_series_label))

        if law_summary_data.business_areas_affected:
            self._add_info_item(
                doc,
                "BUSINESS AREAS AFFECTED",
                str(law_summary_data.business_areas_affected),
            )

        revenue_penalties_from_subject = (
            law_summary_data.subject_matter.revenue_based_penalties or "N/A"
        )
        self._add_info_item(
            doc, "REVENUE-BASED PENALTIES", revenue_penalties_from_subject
        )

        key_stakeholder_roles = (
            law_summary_data.subject_matter.key_stakeholder_roles or "N/A"
        )
        self._add_info_item(doc, "KEY STAKEHOLDER ROLES", key_stakeholder_roles)

        if law_summary_data.source_link:
            self._add_info_item_with_link(
                doc,
                "SOURCE LINK",
                "View Original Document",
                law_summary_data.source_link,
            )

        doc.add_paragraph()  # Spacing

        # SUBJECT MATTER SECTION
        doc.add_paragraph("Subject Matter", style="SectionHeader")
        # Use scope_subject_matter_summary from the structured object
        subject_content = law_summary_data.subject_matter.scope_subject_matter_summary
        if subject_content:
            content = self._process_content(subject_content)
            self._add_content_box(doc, content)
        doc.add_paragraph()  # Spacing

        # ROLES, RESPONSIBILITIES AND PENALTIES SECTION
        doc.add_paragraph(
            "Roles, Responsibilities and Penalties", style="SectionHeader"
        )

        # Revenue-Based Penalties Alert with red background for YES
        revenue_status = law_summary_data.roles_penalties.revenue_based_penalties_status
        if revenue_status:
            alert_text = f"Revenue-Based Penalties: {revenue_status}"
            if revenue_status.upper() == "YES":
                alert_text += "\nThis law includes penalties calculated as a percentage of company revenue/turnover."
                # Use red background for YES status
                self._add_alert_box(doc, alert_text, is_danger=True)
            else:
                # Use normal styling for other statuses
                self._add_alert_box(doc, alert_text, is_danger=False)

        # Roles table
        if law_summary_data.roles_raw:
            self._create_roles_table(doc, law_summary_data.roles_raw)

        # General Penalties subsection
        if (
            law_summary_data.roles_penalties.general_penalties_raw
            and law_summary_data.roles_penalties.general_penalties_raw.strip()
        ):
            doc.add_paragraph(
                "General Penalties Not Role-Specific", style="SubsectionHeader"
            )
            content = self._process_content(
                law_summary_data.roles_penalties.general_penalties_raw
            )
            self._add_content_box(doc, content)

        # Penalty Severity Assessment subsection
        if (
            law_summary_data.roles_penalties.penalty_severity_assessment_raw
            and law_summary_data.roles_penalties.penalty_severity_assessment_raw.strip()
        ):
            doc.add_paragraph("Penalty Severity Assessment", style="SubsectionHeader")
            content = self._process_content(
                law_summary_data.roles_penalties.penalty_severity_assessment_raw
            )
            self._add_content_box(doc, content)

        doc.add_paragraph()  # Spacing

        # TIMELINE FOR COMPLIANCE SECTION (moved to end as per HTML)
        doc.add_paragraph("Timeline for Compliance", style="SectionHeader")
        if law_summary_data.timeline.timeline_content:
            content = self._process_content(law_summary_data.timeline.timeline_content)
            self._add_content_box(doc, content)
        doc.add_paragraph()  # Spacing

        # ADDITIONAL RESOURCES SECTION (if present)
        if law_summary_data.full_report_link:
            doc.add_paragraph("Additional Resources", style="SectionHeader")
            self._add_content_box(doc, "Download Full PDF Report")

        # Convert to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _setup_styles(self, doc: DocxDocument) -> None:
        """Setup professional styles matching the PDF layout."""
        styles = doc.styles

        # Main Title style - same size as PDF
        if "MainTitle" not in [style.name for style in styles]:
            title_style = styles.add_style("MainTitle", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = "Calibri"
            title_style.font.size = Pt(14)  # Same as PDF
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 0, 0)  # Black like PDF
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_style.paragraph_format.space_after = Pt(12)
            title_style.paragraph_format.space_before = Pt(0)

        # Metadata style
        if "Metadata" not in [style.name for style in styles]:
            meta_style = styles.add_style("Metadata", WD_STYLE_TYPE.PARAGRAPH)
            meta_style.font.name = "Calibri"
            meta_style.font.size = Pt(12)  # Same as PDF
            meta_style.font.color.rgb = RGBColor(0, 0, 0)
            meta_style.paragraph_format.space_after = Pt(6)
            meta_style.paragraph_format.space_before = Pt(4)

        # Section header style - same as PDF
        if "SectionHeader" not in [style.name for style in styles]:
            section_style = styles.add_style("SectionHeader", WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = "Calibri"
            section_style.font.size = Pt(14)  # Same as PDF
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0, 0, 0)
            section_style.paragraph_format.space_before = Pt(16)
            section_style.paragraph_format.space_after = Pt(8)

        # Subsection header style
        if "SubsectionHeader" not in [style.name for style in styles]:
            subsection_style = styles.add_style(
                "SubsectionHeader", WD_STYLE_TYPE.PARAGRAPH
            )
            subsection_style.font.name = "Calibri"
            subsection_style.font.size = Pt(14)  # Same as PDF
            subsection_style.font.bold = True
            subsection_style.font.color.rgb = RGBColor(0, 0, 0)
            subsection_style.paragraph_format.space_before = Pt(16)
            subsection_style.paragraph_format.space_after = Pt(8)

        # Content style - reduced spacing
        if "Content" not in [style.name for style in styles]:
            content_style = styles.add_style("Content", WD_STYLE_TYPE.PARAGRAPH)
            content_style.font.name = "Calibri"
            content_style.font.size = Pt(12)  # Same as PDF
            content_style.font.color.rgb = RGBColor(0, 0, 0)
            content_style.paragraph_format.space_after = Pt(6)  # Reduced from 12
            content_style.paragraph_format.space_before = Pt(6)  # Reduced from 12
            content_style.paragraph_format.line_spacing = 1.2  # Reduced from 2.0

        # Info item style
        if "InfoItem" not in [style.name for style in styles]:
            info_style = styles.add_style("InfoItem", WD_STYLE_TYPE.PARAGRAPH)
            info_style.font.name = "Calibri"
            info_style.font.size = Pt(12)
            info_style.font.color.rgb = RGBColor(0, 0, 0)
            info_style.paragraph_format.space_after = Pt(8)

        # Alert style for revenue-based penalties
        if "AlertDanger" not in [style.name for style in styles]:
            alert_style = styles.add_style("AlertDanger", WD_STYLE_TYPE.PARAGRAPH)
            alert_style.font.name = "Calibri"
            alert_style.font.size = Pt(12)
            alert_style.font.color.rgb = RGBColor(239, 68, 68)  # Red text
            alert_style.paragraph_format.space_after = Pt(16)
            alert_style.paragraph_format.space_before = Pt(16)

        if "AlertNormal" not in [style.name for style in styles]:
            alert_normal_style = styles.add_style(
                "AlertNormal", WD_STYLE_TYPE.PARAGRAPH
            )
            alert_normal_style.font.name = "Calibri"
            alert_normal_style.font.size = Pt(12)
            alert_normal_style.font.color.rgb = RGBColor(0, 0, 0)
            alert_normal_style.paragraph_format.space_after = Pt(8)
            alert_normal_style.paragraph_format.space_before = Pt(0)

    def _add_info_item(self, doc: DocxDocument, label: str, value: str) -> None:
        """Add an info item like the PDF layout."""
        p = doc.add_paragraph(style="InfoItem")
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    def _add_info_item_with_link(
        self, doc: DocxDocument, label: str, text: str, url: str
    ) -> None:
        """Add an info item with a hyperlink."""
        p = doc.add_paragraph(style="InfoItem")
        p.add_run(f"{label}: ").bold = True
        self._add_hyperlink(p, text, url)

    def _add_hyperlink(self, paragraph: Paragraph, text: str, url: str) -> None:
        """Add a hyperlink to a paragraph."""
        # Create the hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(
            qn("r:id"),
            paragraph.part.relate_to(
                url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            ),
        )

        # Create a new run object containing the link text
        new_run = OxmlElement("w:r")

        # Create run properties and set the style for hyperlinks
        rPr = OxmlElement("w:rPr")

        # Add blue color for the hyperlink
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)

        # Add underline for the hyperlink
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        new_run.append(rPr)

        # Add the text to the run
        new_run.text = text

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _add_content_box(self, doc: DocxDocument, content: str) -> None:
        """Add content box matching PDF styling with proper bold formatting and line breaks."""
        # Split content by double newlines to create separate paragraphs
        paragraphs = content.split("\n\n")

        for para_content in paragraphs:
            if para_content.strip():
                p = doc.add_paragraph(style="Content")
                self._add_formatted_text(p, para_content.strip())

    def _add_alert_box(
        self, doc: DocxDocument, alert_text: str, is_danger: bool = False
    ) -> None:
        """Add alert box with optional red styling for danger alerts."""
        if is_danger:
            # Add spacing before the alert
            doc.add_paragraph()

            # Add red background effect using table cell
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.cell(0, 0)

            # Set cell background to light red
            cell_xml = cell._tc
            cell_properties = cell_xml.get_or_add_tcPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "fee2e2")  # Light red background
            cell_properties.append(shading)

            # Add text with black color and proper padding
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(12)  # Top padding
            p.paragraph_format.space_after = Pt(12)  # Bottom padding
            p.paragraph_format.left_indent = Pt(12)  # Left padding
            p.paragraph_format.right_indent = Pt(12)  # Right padding

            run = p.add_run(alert_text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black text for maximum contrast
            run.font.size = Pt(12)

            # Minimal cell margins since we're using paragraph formatting for padding
            cell.top_margin = Inches(0.05)
            cell.bottom_margin = Inches(0.05)
            cell.left_margin = Inches(0.05)
            cell.right_margin = Inches(0.05)

            # Add spacing after the alert
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(alert_text, style="AlertNormal")
            p.runs[0].bold = True

    def _create_roles_table(self, doc: DocxDocument, roles_raw: str) -> None:
        """Create roles table matching HTML/PDF structure."""
        if not roles_raw:
            return

        # Parse roles data
        role_rows = []
        for line in roles_raw.split("\n"):
            if "|" in line and line.strip() and not line.startswith("="):
                parts = line.split("|")
                if len(parts) >= 3:
                    role_rows.append(
                        [
                            parts[0].strip(),
                            self._clean_html_content(parts[1].strip()),
                            self._clean_html_content(parts[2].strip()),
                        ]
                    )
                elif len(parts) == 2:
                    role_rows.append(
                        [
                            parts[0].strip(),
                            self._clean_html_content(parts[1].strip()),
                            "",
                        ]
                    )

        if not role_rows:
            return

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Role"
        hdr_cells[1].text = "Responsibilities"
        hdr_cells[2].text = "Penalties"

        # Make headers bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add data rows
        for row_data in role_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = row_data[0]
            row_cells[1].text = row_data[1]
            row_cells[2].text = row_data[2]

    def _process_content(self, content: str) -> str:
        """Process content by transforming markers to formatting for Word."""
        if not content:
            return ""

        # Transform =text= markers and clean HTML
        content = self._transform_markers_to_bold(content)
        content = self._clean_html_content(content)

        return content

    def _transform_markers_to_bold(self, text: str) -> str:
        """Transform =text= markers to bold section titles with proper spacing."""
        if not text:
            return ""

        # Convert =Section= markers to bold text markers with double newlines for separation
        text = re.sub(r"=([^=]+)=", r"\n\n**\1**\n", text)

        # Clean up multiple consecutive newlines at the start
        text = re.sub(r"^\n+", "", text)

        return text

    def _clean_html_content(self, text: str) -> str:
        """Clean HTML tags from text for Word formatting."""
        if not text:
            return ""

        # Convert <br/> tags to newlines
        text = text.replace("<br/>", "\n").replace("<br>", "\n")

        # Remove HTML tags but preserve bold markers
        text = re.sub(r"<b>(.*?)</b>", r"**\1**", text, flags=re.IGNORECASE)
        text = re.sub(r"<strong>(.*?)</strong>", r"**\1**", text, flags=re.IGNORECASE)

        # Remove other HTML tags
        text = re.sub(r"<[^>]+>", "", text)

        # Clean up whitespace
        text = re.sub(r"\n\s*\n", "\n\n", text)
        text = text.strip()

        return text

    def _add_formatted_text(self, paragraph: Paragraph, text: str) -> None:
        """Add text to paragraph with proper bold formatting for Word."""
        if not text:
            return

        # Split text by **bold** markers and process
        parts = re.split(r"\*\*(.*?)\*\*", text)

        for i, part in enumerate(parts):
            if not part:
                continue

            # Even indices are normal text, odd indices are bold
            if i % 2 == 0:
                # Normal text - split by newlines to handle line breaks
                lines = part.split("\n")
                for j, line in enumerate(lines):
                    if line:
                        paragraph.add_run(line)
                    if j < len(lines) - 1:  # Add line break if not last line
                        paragraph.add_run("\n")
            else:
                # Bold text
                run = paragraph.add_run(part)
                run.bold = True
